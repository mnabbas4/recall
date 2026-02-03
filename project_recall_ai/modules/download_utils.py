import pandas as pd
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document


# ===============================
# CSV
# ===============================
def export_csv(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    buffer.write(f"# SUMMARY\n# {summary.replace('\n', ' ')}\n\n".encode())
    df.to_csv(buffer, index=False)
    return buffer.getvalue()


# ===============================
# EXCEL
# ===============================
def export_excel(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Results", index=False)

        summary_df = pd.DataFrame({"Summary": summary.split("\n")})
        summary_df.to_excel(writer, sheet_name="Summary", index=False)

    return buffer.getvalue()


# ===============================
# PDF
# ===============================
def export_pdf(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "AI Query Summary")
    y -= 30

    c.setFont("Helvetica", 10)
    for line in summary.split("\n"):
        c.drawString(40, y, line)
        y -= 14
        if y < 60:
            c.showPage()
            y = height - 40

    c.showPage()
    y = height - 40
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Results Table")
    y -= 20

    c.setFont("Helvetica", 8)
    for _, row in df.head(40).iterrows():  # safety limit
        row_text = " | ".join(str(v)[:30] for v in row.values)
        c.drawString(40, y, row_text)
        y -= 12
        if y < 40:
            c.showPage()
            y = height - 40

    c.save()
    return buffer.getvalue()


# ===============================
# WORD
# ===============================
def export_word(df: pd.DataFrame, summary: str) -> bytes:
    doc = Document()
    doc.add_heading("AI Query Summary", level=1)

    for line in summary.split("\n"):
        doc.add_paragraph(line)

    doc.add_heading("Results Table", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdrs = table.rows[0].cells

    for i, col in enumerate(df.columns):
        hdrs[i].text = col

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
